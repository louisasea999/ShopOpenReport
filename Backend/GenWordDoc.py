# -*- coding: utf-8 -*-
# @Time    : 18-10-10 下午10:57
# @Author  : Yanming
# @Email   : 
# @File    : Test4Word.py
# @Software: PyCharm

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from Utils import createTable,createStyles,createMainTitle
from docx.shared import RGBColor
from docx.enum.text import WD_LINE_SPACING


document = Document()

# ID = 01    -----------------------------------------------------------------------------------------

#空一行
paragraph = document.add_paragraph(u'\n')
#标题的大小，传给createMainTitle
ptSize = 26
#基于ID01 标题  jason children 参数
customer = 'yum'
market = '市场'
code = '新店代码'
date = '日期'
report = '新店报告'
#拼接所有的参数
headerTitle = customer+market+code+date+report
#创建自定义style
createStyles(document)
#创建标题
createMainTitle(headerTitle, document,ptSize)
#空两行
paragraph = document.add_paragraph(u'\n')
paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
# ID = 01    -----------------------------------------------------------------------------------------

# ID = 02    -----------------------------------------------------------------------------------------
#基于ID02 标题  jason children 参数
version = 'Version0.1'
#标题的大小，传给createMainTitle
ptSize = 14
#创建标题
createMainTitle(version, document,ptSize)
#换页
document.add_page_break()
# ID = 02    -----------------------------------------------------------------------------------------

# ID = 03    -----------------------------------------------------------------------------------------
#空一行
paragraph = document.add_paragraph(u'\n')
#基于ID03 标题  jason children 参数
preface = '引言'
#标题的大小，传给createMainTitle
ptSize = 21
#创建标题
createMainTitle(preface, document,ptSize)
#换页
document.add_page_break()
# ID = 03    -----------------------------------------------------------------------------------------

# ID = 04    -----------------------------------------------------------------------------------------
#基于ID04 标题  jason children 参数
tableName_ID04 = '新店基础信息表'
titleLevel = 1
rowNum = 21
colNum = 2
colList = ['餐厅编号','餐厅名称','餐厅电话','品牌','开业日期','IT调试工程师品牌','POS数量',
           '票打数量','摄像头数量','宽带线路数','KDS数量','无线AP、移动NC',
           'NC数量','DMB TV数','DMB NC','餐厅地址','宽带账号(Mail)','宽带密码(Mail)','宽带账号(DMB)','宽带账号(外送)',
           '宽带密码(外送)']
#添加level1小标题，创建两列的表格
createTable(document,tableName_ID04, titleLevel, rowNum, colNum, colList)
#换页
document.add_page_break()
# ID = 04    -----------------------------------------------------------------------------------------

# ID = 05    -----------------------------------------------------------------------------------------
#基于ID05 标题  jason children 参数
tableName_ID05 = '有效安装时间统计表（小时）'
titleLevel = 1
rowNum = 10
colNum = 2
colList = ['餐厅','总时间','CCTV','POS','POS打印机','PC/Server','NC','Soft','UPS','Other']
#添加level1小标题，创建两列的表格
createTable(document,tableName_ID05, titleLevel, rowNum, colNum, colList)
#换页
document.add_page_break()
# ID = 05    -----------------------------------------------------------------------------------------

# ID = 06    -----------------------------------------------------------------------------------------
#基于ID06 标题  jason children 参数
tableName_ID06 = '标题验收单据（安装单据，Standby单据）'
titleLevel = 1
rowNum = 2
colNum = 2
colList = ['安装单据','Standby单据']
#添加level1小标题，创建两列的表格
createTable(document,tableName_ID06, titleLevel, rowNum, colNum, colList)
#换页
document.add_page_break()
# ID = 06    -----------------------------------------------------------------------------------------

# ID = 07    -----------------------------------------------------------------------------------------
#基于ID07 标题  jason children 参数
tableName_ID07 = '交接表（Vendor、营建、营运）'
titleLevel = 1
rowNum = 1
colNum = 2
colList = ['交接表（Vendor、营建、营运）']
#添加level1小标题，创建两列的表格
createTable(document,tableName_ID07, titleLevel, rowNum, colNum, colList)
#换页
document.add_page_break()
# ID = 07    -----------------------------------------------------------------------------------------

# ID = 08    -----------------------------------------------------------------------------------------
#基于ID08 标题  jason children 参数
tableName_ID08 = '硬件资产收集'
titleLevel = 1
rowNum = 1
colNum = 2
colList = ['硬件资产收集']
#添加level1小标题，创建两列的表格
createTable(document,tableName_ID08, titleLevel, rowNum, colNum, colList)
#换页
document.add_page_break()
# ID = 08    -----------------------------------------------------------------------------------------

# ID = 09    -----------------------------------------------------------------------------------------
#ID09第一标题
docTitle09 = '新店现场照片'
#不使用工具类，直接创建，level=1
T09 = document.add_heading(docTitle09, 1)
T09.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
#ID0901第二级标题
docTitle0901 = '经理室机柜走线及布局'
#不使用工具类，直接创建，level=2
T0901 = document.add_heading(docTitle0901, 2)
T0901.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
#基于ID090101 标题  jason children 参数
tableName_ID090101 = '机柜安装前线路整理、捆扎照片'
titleLevel = 3
rowNum = 1
colNum = 2
colList = ['要求：\n1、全景照片：机柜位置上至天花板下至底平面。\n2、两束线捆扎规范整齐\n3、电源面板规范（两红一白）']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID090101, titleLevel, rowNum, colNum, colList)

#基于ID090102 标题  jason children 参数
tableName_ID090102 = '机柜进出线'
titleLevel = 3
rowNum = 1
colNum = 2
colList = ['机柜进、出线要求:']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID090102, titleLevel, rowNum, colNum, colList)

#基于ID090103 标题  jason children 参数
tableName_ID090103 = '机柜安装前线路整理、捆扎照片'
titleLevel = 3
rowNum = 2
colNum = 2
colList = ['机柜未装设备照片要求：\n1、内部线路走向清晰，能看清楚其基本走向。\n2、网线、视频线辨识清晰（网线前部捆扎，视频线后部捆扎。',
           '机柜安装空间要求:\n1.700(宽)X650(深)X1400(高) (mm)\n2.机柜安装位置配电要求：两个UPS供电的250V/10A三极电源插座，一个市电250V/10A三极电源插座，插座底边离地面高度为1200mm']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID090103, titleLevel, rowNum, colNum, colList)

#基于ID090104 标题  jason children 参数
tableName_ID090104 = '机柜设备安装后开门照'
titleLevel = 3
rowNum = 2
colNum = 2
colList = ['机柜设备开门照片要求：\n1、设备安装整齐，标准。\n2、无飞线、绕线现象。\n3、标识张贴规范清楚\n4、下列机柜示意图为落地机柜，餐厅如为架空机柜，请参照架空机柜标准。',
           '机柜内设备\n1、排线原则：所有网线、信号线、电源线都排在机柜后部，机柜前面不得有飞线\n2、需要将冗余的直流电源线绕成团并用尼龙扎带扎好\n3、联到Switch的网线从Switch上面走线，连接好后用尼龙扎带扎好\n4、H3C设备上需要按用途分别贴“邮件线路”或“外送线路”标签\n5、网线、电话线、视频线需要贴标签']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID090104, titleLevel, rowNum, colNum, colList)

#基于ID090105 标题  jason children 参数
tableName_ID090105 = '交换机、防火墙、宽带猫特写'
titleLevel = 3
rowNum = 1
colNum = 2
colList = ['交换机网线照片要求：\n1、架空式机柜第1层放交换机、第2层放防火墙、第3层托盘放宽带猫（落地式方向相反）\n2、网线端横向拉线固定，网线标号清晰。\n3、交换机、防火墙、宽带猫清晰，线路整齐。\n4、邮件防火墙CONSOLE插口连接CONSOLE线']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID090105, titleLevel, rowNum, colNum, colList)

#基于ID090106 标题  jason children 参数
tableName_ID090106 = 'CCTV主机特写照'
titleLevel = 3
rowNum = 1
colNum = 2
colList = ['CCTV主机照片要求：\n1、正面拍到设备安装位置和走线；背面拍到视频线连接在视频卡上的情况。\n2、后端线路接入规范。']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID090106, titleLevel, rowNum, colNum, colList)

#基于ID090107 标题  jason children 参数
tableName_ID090107 = '机柜关门照'
titleLevel = 3
rowNum = 1
colNum = 2
colList = ['机柜设备关门照片要求：\n1、关闭机柜门，锁好门，照片看到外部环境和线路走入情况。']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID090107, titleLevel, rowNum, colNum, colList)
#换页
document.add_page_break()

#ID0902第二级标题
docTitle0902 = '经理室'
#不使用工具类，直接创建，level=2
T0901 = document.add_heading(docTitle0902, 2)
T0901.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

#基于ID090201 标题  jason children 参数
tableName_ID090201 = 'NC侧面照'
titleLevel = 3
rowNum = 1
colNum = 2
colList = ['NC侧面照片要求：\n1、显示器、键盘、鼠标、主机及连线、电源插座、后面的线路摆放是否整洁、插座。']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID090201, titleLevel, rowNum, colNum, colList)

#基于ID090202 标题  jason children 参数
tableName_ID090202 = '后台打印机'
titleLevel = 3
rowNum = 1
colNum = 2
colList = ['后台打印机要求：\n1、要拍到 电源线、数据线走线和电源、网络插座，激光打印机不超过全图五分之一。\n2、打印机如果是网络打印机，则必须使用网线连接，不可使用USB。\n3、打印机不可接入UPS电源。']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID090202, titleLevel, rowNum, colNum, colList)

#基于ID090203 标题  jason children 参数
tableName_ID090203 = '无线AP'
titleLevel = 3
rowNum = 1
colNum = 2
colList = ['1.AP定位标准:营运选定外场AP常用区域,就近指定区域设备检修口\n2.营建供应商负责将二根网线由机柜布到指定AP位置，AP电源放置在配电房\n3.营建供应商负责AP天线打孔,直径:18mm\n4.IT服务商负责AP设备配置与安装，安装空间要求:94X245X245mm']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID090203, titleLevel, rowNum, colNum, colList)

#基于ID090204 标题  jason children 参数
tableName_ID090204 = '移动NC'
titleLevel = 3
rowNum = 1
colNum = 2
colList = ['1.正面开机状态。\n2.要能看清已连接无线信号登录操作平台']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID090204, titleLevel, rowNum, colNum, colList)

#基于ID090205 标题  jason children 参数
tableName_ID090205 = '卡钟整体照'
titleLevel = 3
rowNum = 1
colNum = 2
colList = ['卡钟要求：\n1、要拍到卡钟及周边环境，卡钟在照片正中，不超过全图5分之一大小。\n2、线路规范，隐蔽。\n3、考勤机安装位置由市场营运部门(餐厅经理)确定']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID090205, titleLevel, rowNum, colNum, colList)
#换页
document.add_page_break()

#ID0903第二级标题
docTitle0903 = '收银系统'
#不使用工具类，直接创建，level=2
T0903 = document.add_heading(docTitle0903, 2)
T0903.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

#ID090301第三级标题
docTitle090301 = '柜台'
#不使用工具类，直接创建，level=3
T090301 = document.add_heading(docTitle090301, 3)
T090301.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

#基于ID09030101 标题  jason children 参数
tableName_ID09030101 = 'POS客区照'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['POS1:\n开机收银界面、旁边POS打印机、下面抽屉。\n冗余线路单独分别捆扎，置于POS机正后方。\n网络端口标号清晰且与点位图一致。']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID09030101, titleLevel, rowNum, colNum, colList)

#基于ID09030102 标题  jason children 参数
tableName_ID09030102 = 'POS员工区照'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['POS2\n开机收银界面、旁边POS打印机、下面抽屉。\n冗余线路单独分别捆扎，置于POS机正后方。\n网络端口标号清晰且与点位图一致。']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID09030102, titleLevel, rowNum, colNum, colList)

#基于ID09030103 标题  jason children 参数
tableName_ID09030103 = '单台POS整体照'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['POS2\n开机收银界面、旁边POS打印机、下面抽屉。\n冗余线路单独分别捆扎，置于POS机正后方。\n网络端口标号清晰且与点位图一致。']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID09030103, titleLevel, rowNum, colNum, colList)

#基于ID09030104 标题  jason children 参数
tableName_ID09030104 = '票打整体照'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['POS2\n开机收银界面、旁边POS打印机、下面抽屉。\n冗余线路单独分别捆扎，置于POS机正后方。\n网络端口标号清晰且与点位图一致。']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID09030104, titleLevel, rowNum, colNum, colList)

#基于ID09030105 标题  jason children 参数
tableName_ID09030105 = 'POS抽屉、线路、模块'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['POS2\n开机收银界面、旁边POS打印机、下面抽屉。\n冗余线路单独分别捆扎，置于POS机正后方。\n网络端口标号清晰且与点位图一致。']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID09030105, titleLevel, rowNum, colNum, colList)
#换页
document.add_page_break()

#ID090302第三级标题
docTitle090302 = '点配分离：厨房版本号（请选择模式）柜台'
#不使用工具类，直接创建，level=3
T090302 = document.add_heading(docTitle090302, 3)
T090302.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

#基于ID09030201 标题  jason children 参数
tableName_ID09030201 = 'KDS正面照-配餐屏'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1.冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2.网络模块标号清晰且与点位图一致\n3.测试样张图片；无飞线\n4.对应打印机的测试小票']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID09030201, titleLevel, rowNum, colNum, colList)

#基于ID09030202 标题  jason children 参数
tableName_ID09030202 = 'KDS侧面照-配餐'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1.冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2.网络模块标号清晰且与点位图一致\n3.测试样张图片；无飞线\n4.对应打印机的测试小票']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID09030202, titleLevel, rowNum, colNum, colList)

#基于ID09030203 标题  jason children 参数
tableName_ID09030203 = 'KDS正面照-汉堡屏'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1.冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2.网络模块标号清晰且与点位图一致\n3.测试样张图片；无飞线\n4.对应打印机的测试小票']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID09030203, titleLevel, rowNum, colNum, colList)

#基于ID09030204 标题  jason children 参数
tableName_ID09030204 = 'KDS侧面照-汉堡屏'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1.冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2.网络模块标号清晰且与点位图一致\n3.测试样张图片；无飞线\n4.对应打印机的测试小票']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID09030204, titleLevel, rowNum, colNum, colList)

#基于ID09030205 标题  jason children 参数
tableName_ID09030205 = 'KDS正面照-饮料屏'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1.冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2.网络模块标号清晰且与点位图一致\n3.测试样张图片；无飞线\n4.对应打印机的测试小票']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID09030205, titleLevel, rowNum, colNum, colList)

#基于ID09030206 标题  jason children 参数
tableName_ID09030206 = 'KDS侧面照-饮料屏'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1.冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2.网络模块标号清晰且与点位图一致\n3.测试样张图片；无飞线\n4.对应打印机的测试小票']
#添加level3小标题，创建两列的表格
createTable(document,tableName_ID09030206, titleLevel, rowNum, colNum, colList)
#换页
document.add_page_break()

#ID090303第三级标题
docTitle090303 = '取餐屏'
#不使用工具类，直接创建，level=3
T090303 = document.add_heading(docTitle090303, 3)
T090303.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

#基于ID09030301 标题  jason children 参数
tableName_ID09030301 = '取餐屏'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID09030301, titleLevel, rowNum, colNum, colList)
#换页
document.add_page_break()

#ID090304第三级标题
docTitle090304 = '2GBOH厨房标准图及参数描述'
#不使用工具类，直接创建，level=3
T090304 = document.add_heading(docTitle090304, 3)
T090304.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

#基于ID09030401 标题  jason children 参数
tableName_ID09030401 = '2GBOH厨房标准图及参数描述'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID09030401, titleLevel, rowNum, colNum, colList)

#基于ID09030402 标题  jason children 参数
tableName_ID09030402 = '虚机网络配置截图'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID09030402, titleLevel, rowNum, colNum, colList)

#基于ID09030403 标题  jason children 参数
tableName_ID09030403 = '厨房设备全景图'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID09030403, titleLevel, rowNum, colNum, colList)

#基于ID09030404 标题  jason children 参数
tableName_ID09030404 = '总配：打印机定位原则'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID09030404, titleLevel, rowNum, colNum, colList)
#换页
document.add_page_break()

#ID0904第二级标题
docTitle0904 = '周边收银系统'
#不使用工具类，直接创建，level=2
T0904 = document.add_heading(docTitle0904, 2)
T0904.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

#ID090401第三级标题
docTitle090401 = '外送：外送POS整体照'
#不使用工具类，直接创建，level=3
T090401 = document.add_heading(docTitle090401, 3)
T090401.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

#基于ID09040101 标题  jason children 参数
tableName_ID09040101 = '外送：外送POS整体照'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID09040101, titleLevel, rowNum, colNum, colList)

#ID090402第三级标题
docTitle090402 = '甜品站'
#不使用工具类，直接创建，level=3
T090402 = document.add_heading(docTitle090402, 3)
T090402.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

#基于ID09040201 标题  jason children 参数
tableName_ID09040201 = '单台POS整体照'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID09040201, titleLevel, rowNum, colNum, colList)

#基于ID09040202 标题  jason children 参数
tableName_ID09040202 = '票打整体照'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID09040202, titleLevel, rowNum, colNum, colList)

#基于ID09040203 标题  jason children 参数
tableName_ID09040203 = 'POS抽屉、线路、模块'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID09040203, titleLevel, rowNum, colNum, colList)
#换页
document.add_page_break()

#ID090403第三级标题
docTitle090403 = '外卖点'
#不使用工具类，直接创建，level=3
T090403 = document.add_heading(docTitle090403, 3)
T090403.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

#基于ID09040301 标题  jason children 参数
tableName_ID09040301 = '单台POS整体照'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID09040301, titleLevel, rowNum, colNum, colList)

#基于ID09040302 标题  jason children 参数
tableName_ID09040302 = '票打整体照'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID09040302, titleLevel, rowNum, colNum, colList)

#基于ID09040303 标题  jason children 参数
tableName_ID09040303 = 'POS抽屉、线路、模块'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID09040303, titleLevel, rowNum, colNum, colList)
#换页
document.add_page_break()

#ID090404第三级标题
docTitle090404 = '车道POS机'
#不使用工具类，直接创建，level=3
T090404 = document.add_heading(docTitle090404, 3)
T090404.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

#基于ID09040401 标题  jason children 参数
tableName_ID09040401 = '单台POS整体照'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID09040401, titleLevel, rowNum, colNum, colList)

#基于ID09040402 标题  jason children 参数
tableName_ID09040402 = '票打整体照'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID09040402, titleLevel, rowNum, colNum, colList)

#基于ID09040403 标题  jason children 参数
tableName_ID09040403 = 'POS抽屉、线路、模块'
titleLevel = 4
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID09040403, titleLevel, rowNum, colNum, colList)
#换页
document.add_page_break()

#ID0905第二级标题
docTitle0905 = 'DMB电子餐牌设置'
#不使用工具类，直接创建，level=2
T0905 = document.add_heading(docTitle0905, 2)
T0905.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

#基于ID090501 标题  jason children 参数
tableName_ID090501 = 'POS抽屉、线路、模块'
titleLevel = 3
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID090501, titleLevel, rowNum, colNum, colList)

#基于ID090502 标题  jason children 参数
tableName_ID090502 = 'DMB全部屏显播放正常的照片'
titleLevel = 3
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID090502, titleLevel, rowNum, colNum, colList)

#基于ID090503 标题  jason children 参数
tableName_ID090503 = 'DMB标签'
titleLevel = 3
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID090503, titleLevel, rowNum, colNum, colList)

#ID0906第二级标题
docTitle0906 = 'CCTV监控'
#不使用工具类，直接创建，level=2
T0906 = document.add_heading(docTitle0906, 2)
T0906.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

#基于ID090601 标题  jason children 参数
tableName_ID090601 = 'CCTV监控画面'
titleLevel = 3
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID090601, titleLevel, rowNum, colNum, colList)
#换页
document.add_page_break()

#ID0907第二级标题
docTitle0907 = '其他'
#不使用工具类，直接创建，level=2
T0907 = document.add_heading(docTitle0907, 2)
T0907.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

#基于ID090701 标题  jason children 参数
tableName_ID090701 = 'SDCD'
titleLevel = 3
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID090701, titleLevel, rowNum, colNum, colList)

#基于ID090702 标题  jason children 参数
tableName_ID090702 = 'TVBOX'
titleLevel = 3
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID090702, titleLevel, rowNum, colNum, colList)

#基于ID090703 标题  jason children 参数
tableName_ID090703 = 'UPS：UPS正面照'
titleLevel = 3
rowNum = 1
colNum = 2
colList = ['1、冗余的电源线、数据线、网线需卷好，并用尼龙扎带扎好置于工作台后侧\n2、网络模块标号清晰且与点位图一致\n3、测试样张图片，无飞线']
#添加level4小标题，创建两列的表格
createTable(document,tableName_ID090703, titleLevel, rowNum, colNum, colList)
#换页
document.add_page_break()
# ID = 09    -----------------------------------------------------------------------------------------

# ID = 10    -----------------------------------------------------------------------------------------
#ID10第一标题
docTitle10 = '新店功能性验收'
#不使用工具类，直接创建，level=1
T10 = document.add_heading(docTitle10, 1)
T10.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

#基于ID1001 标题  jason children 参数
tableName_ID1001 = '每台POS的Ready打印单，收银机与打印机工作正常'
titleLevel = 2
rowNum = 1
colNum = 2
colList = ['每台POS的Ready打印单，收银机与打印机工作正常']
#添加level2小标题，创建两列的表格
createTable(document,tableName_ID1001, titleLevel, rowNum, colNum, colList)

#基于ID1002 标题  jason children 参数
tableName_ID1002 = 'Compris数据清零后，打印Sales Report'
titleLevel = 2
rowNum = 1
colNum = 2
colList = ['Compris数据清零后，打印Sales Report']
#添加level2小标题，创建两列的表格
createTable(document,tableName_ID1002, titleLevel, rowNum, colNum, colList)

#基于ID1003 标题  jason children 参数
tableName_ID1003 = '打印测试发票'
titleLevel = 2
rowNum = 1
colNum = 2
colList = ['打印测试发票']
#添加level2小标题，创建两列的表格
createTable(document,tableName_ID1003, titleLevel, rowNum, colNum, colList)

#基于ID1004 标题  jason children 参数
tableName_ID1004 = '打印外送测试单'
titleLevel = 2
rowNum = 1
colNum = 2
colList = ['打印外送测试单']
#添加level2小标题，创建两列的表格
createTable(document,tableName_ID1004, titleLevel, rowNum, colNum, colList)

#基于ID1005 标题  jason children 参数
tableName_ID1005 = '餐条显示正确的门店编号、门店中文地址、门店电话'
titleLevel = 2
rowNum = 1
colNum = 2
colList = ['餐条显示正确的门店编号、门店中文地址、门店电话']
#添加level2小标题，创建两列的表格
createTable(document,tableName_ID1005, titleLevel, rowNum, colNum, colList)
#换页
document.add_page_break()
# ID = 10    -----------------------------------------------------------------------------------------

# ID = 11    -----------------------------------------------------------------------------------------
#基于ID11 标题  jason children 参数
tableName_ID11 = '新店安装过程问题汇总'
titleLevel = 2
rowNum = 6
colNum = 2
colList = ['门店代码','营建方面问题','门店方面问题','设备方面问题','线路方面问题','工程师现场未完成/无法完成/等待时间超过2小时等事宜&建议']
#添加level2小标题，创建两列的表格
table01 = createTable(document,tableName_ID11, titleLevel, rowNum, colNum, colList)

# widths = (Inches(1), Inches(6))
# for row in table01.rows:
#     for idx, width in enumerate(widths):
#         row.cells[idx].width = width

#换页
document.add_page_break()
# ID = 11    -----------------------------------------------------------------------------------------

document.save('demo1.docx')  # 保存这个文档