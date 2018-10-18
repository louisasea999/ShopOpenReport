
# -*- coding: utf-8 -*-
# @Time    : 18-10-12 上午9:54
# @Author  : Yanming
# @Email   : 
# @File    : Utils.py
# @Software: PyCharm

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.shared import Cm



def createTable(document,titleName, titleLevel, rowNum, colNum, colList):
    # table = document.add_table(rows=10, cols=2, style="ColorfulList-Accent5")
    # set_column_width(table.columns[1], Cm(13))
    docTitle = titleName  # add title which ID equal jason file 01
    T01 = document.add_heading(docTitle, titleLevel)  # add_heading level is 0.   first level title
    T01.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # make 01 title align center
    table01 = document.add_table(rows=rowNum, cols=colNum, style='Table Grid')  # 添加一个表格
    hdr_rows = table01.rows
    print(len(hdr_rows))
    for hdr_row in range(len(hdr_rows)):
        hdr_cells = table01.rows[hdr_row].cells
        hdr_cells[0].text = colList[hdr_row]
    document.save('demoTestUtils.docx')  # 保存这个文档

    return table01

def createStyles(document):
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('HeadingStyle',WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(10)
    obj_font.name = '微软雅黑'
    obj_font.color.rgb = RGBColor(0x71, 0x71, 0xc6)

    obj_charstyle = obj_styles.add_style('ReportTitleStyle', WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(26)
    obj_font.name = '宋体'
    obj_font.color.rgb = RGBColor(0x12, 0x12, 0x12)
    return obj_charstyle

def createMainTitle(mainTitle, document,ptSize):
    document.styles['Normal'].font.name = u'宋体'  # 可换成word里面任意字体
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中设置
    run = p.add_run(mainTitle)
    run.font.color.rgb = RGBColor(12, 12, 12)  # 颜色设置，这里是用RGB颜色
    run.font.size = Pt(ptSize)  # 字体大小设置，和word里面的字号相对应
    run.bold = True

def set_column_width(column, width):
    for cell in column.cells:
        cell.width = width





